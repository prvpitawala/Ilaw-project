import re

def format_text(text):
    # Replace **bold** with <b>bold</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Replace newline characters \n with <br> tags
    text = text.replace('\n', '<br>')
   
    # Replace * with <li>
    text = text.replace('* ', '<li>').replace('\n', '</li>\n')
    
    # Wrap the whole package descriptions in <ul> tags for better HTML structure
    #text = '<ul>\n' + text + '</li>\n</ul>'

    while text.endswith('<br>'):
        text = text[:-4] 
        
    return text

#########################################################################################
#stage2 impliments
############################################################################################
# Load environment variables from .env file
import os
import docx
from odf import text, teletype
from odf.opendocument import load
from typing import Dict
from typing import List, Dict, Any, Tuple
from tqdm import tqdm
import openai
from dotenv import load_dotenv
import chromadb
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from docx import Document  # python-docx for .docx files
from pdfminer.high_level import extract_text  # Keeping existing PDF support

def get_app_data_dir():
    app_data_dir = os.path.join(os.environ['APPDATA'], "Rag Doc System")
    if not os.path.exists(app_data_dir):
        os.makedirs(app_data_dir)
    return app_data_dir

def fileFolderPathGen(relativePath):
    return os.path.join(get_app_data_dir(), relativePath)

ENV_DIR = fileFolderPathGen('.env') #r"../.env"
load_dotenv(dotenv_path=ENV_DIR)

# Configure OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# def get_documents_from_folder(folder_path: str) -> Dict[str, str]:
#     """
#     Extract text documents from a folder
    
#     Args:
#         folder_path: Path to the folder containing documents
        
#     Returns:
#         Dictionary mapping filenames to their content
#     """
#     documents = {}
    
#     # Process each file in the folder and its subfolders
#     for root, _, files in os.walk(folder_path):
#         for file in files:
#             # Focus on text files
#             if file.endswith(('.txt')):
#                 file_path = os.path.join(root, file)
#                 try:
#                     with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
#                         content = f.read()
#                         relative_path = os.path.relpath(file_path, folder_path)
#                         documents[relative_path] = content
#                 except Exception as e:
#                     print(f"Error reading file {file_path}: {e}")

#             # Focus on pdf files 
#             if file.endswith(('.pdf')):
#                 file_path = os.path.join(root, file)
#                 try:
#                     content = extract_text(file_path)
#                     relative_path = os.path.relpath(file_path, folder_path)
#                     txt_filename = f"{os.path.splitext(relative_path)[0]}.txt"

#                     # save in the document dic as .txt
#                     documents[txt_filename] = content
#                 except Exception as e:
#                     print(f"Error reading file {file_path}: {e}")
    
#     return documents

def get_documents_from_folder(folder_path: str) -> Dict[str, str]:
    """
    Extract text documents from a folder
    
    Args:
        folder_path: Path to the folder containing documents
        
    Returns:
        Dictionary mapping filenames to their content
    """
    documents = {}
    
    # Process each file in the folder and its subfolders
    for root, _, files in os.walk(folder_path):
        for file_ in files:
            file_path = os.path.join(root, file_)
            relative_path = os.path.relpath(file_path, folder_path)
            txt_filename = f"{os.path.splitext(relative_path)[0]}.txt"
            
            try:
                # Focus on text files
                if file_.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read()
                        documents[relative_path] = content

                # Focus on pdf files 
                elif file_.endswith('.pdf'):
                    content = extract_text(file_path)
                    documents[txt_filename] = content
                
                # Handle docx files
                elif file_.endswith('.docx'):
                    doc = docx.Document(file_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    documents[txt_filename] = content
                
                # Handle odt files
                elif file_.endswith('.odt'):
                    doc = load(file_path)
                    content = teletype.extractText(doc)
                    documents[txt_filename] = content
                    
            except Exception as e:
                print(f"Error reading file {file_path}: {e}")
    
    return documents   


def get_documents_from_filenames(root: str, filenames: List[str], collection_name: str) -> Dict[str, str]:
    """
    Extract text documents from specific files in a folder
    
    Args:
        root: Path to the root folder
        filenames: List of file names to extract
        collection_name: Name of the collection folder
        
    Returns:
        Dictionary mapping filenames to their content
    """
    folder_path = os.path.join(root, collection_name)
    documents = {}
    
    for filename in filenames:
        file_path = os.path.join(folder_path, filename)
        
        try:
            # Handle text files
            if filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                    documents[filename] = content
            
            # Handle PDF files
            elif filename.endswith('.pdf'):
                content = extract_text(file_path)
                txt_filename = f"{os.path.splitext(filename)[0]}.txt"
                documents[txt_filename] = content
            
            # Handle Word (.docx) files
            elif filename.endswith('.docx'):
                doc = Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                txt_filename = f"{os.path.splitext(filename)[0]}.txt"
                documents[txt_filename] = content
            
            # Handle OpenDocument (.odt) files
            elif filename.endswith('.odt'):
                textdoc = load(file_path)
                allparas = textdoc.getElementsByType(text.P)
                content = '\n'.join([teletype.extractText(para) for para in allparas])
                txt_filename = f"{os.path.splitext(filename)[0]}.txt"
                documents[txt_filename] = content
            
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
    
    return documents


def save_documents_as_text_files_in_folder(documents, output_dir="extracted_documents"):
    """
    Given a dictionary of documents, save each document as a text file in the output directory.
    
    Args:
        documents (dict): Dictionary where keys are filenames and values are file content.
        output_dir (str): The directory where the text files will be saved.
    """
    # Create the output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Iterate over the dictionary
    for filename, content in documents.items():
        # Get the full path for the output text file
        output_path = os.path.join(output_dir, filename)
        
        # Create the necessary directories if they don't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        try:
            # Open the file in write mode and save the content
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"Saved: {output_path}")
        except Exception as e:
            print(f"Error saving {filename}: {e}")




def chunk_text(chunk_size, chunk_overlap, documents: Dict[str, str]) -> List[Dict[str, Any]]:
    """
    Split documents into chunks suitable for embedding
    
    Args:
        documents: Dictionary mapping document names to their content
        
    Returns:
        List of document chunks with metadata
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    
    chunks = []
    for doc_name, content in documents.items():
        splits = text_splitter.split_text(content)
        
        for i, chunk_text in enumerate(splits):
            chunks.append({
                "id": f"{doc_name}_chunk_{i}",
                "text": chunk_text,
                "metadata": {
                    "source": doc_name,
                    "chunk": i
                }
            })
    return chunks



def save_chunks_as_text_files(chunks: List[Dict[str, Any]], output_dir: str = "chunks") -> None:
    """
    Save each document chunk as a separate text file
    
    Args:
        chunks: List of document chunks with metadata
        output_dir: Directory to save the chunk files
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"Saving {len(chunks)} chunks to {output_dir}...")
    
    for chunk in chunks:
        # Create a filename based on the chunk ID
        safe_id = chunk['id'].replace('/', '_').replace('\\', '_')
        filename = f"{safe_id}.txt"

        file_path = os.path.join(output_dir, filename)
        
        # Write the chunk content to a file
        with open(file_path, 'w', encoding='utf-8') as f:
            # Add metadata as a header
            f.write(f"Source: {chunk['metadata']['source']}\n")
            f.write(f"Chunk: {chunk['metadata']['chunk']}\n")
            f.write(f"ID: {chunk['id']}\n")
            f.write("-" * 50 + "\n\n")
            
            # Write the actual chunk text
            f.write(chunk['text'])
    
    print(f"Successfully saved {len(chunks)} chunk files to {output_dir}")


def clear_chroma_database(chroma_client,collection_name=None):
    """
    Clear data from ChromaDB
    
    Args:
        collection_name: Name of the collection to clear. If None, all collections will be deleted.
    """
    
    try:
        if collection_name:
            # Delete a specific collection
            try:
                chroma_client.delete_collection(collection_name)
                print(f"Collection '{collection_name}' has been deleted")
            except ValueError as e:
                print(f"Collection '{collection_name}' has been not deleted")
                return
        else:
            # Get all collections
            all_collections = chroma_client.list_collections()
            
            # Delete each collection
            for collection in all_collections:
                chroma_client.delete_collection(collection.name)
                print(f"Deleted collection: {collection.name}")
            
            print(f"All collections have been deleted")
    except Exception as e:
        print(f"Error clearing database: {e}")


# def get_chatgpt_response(prompt, system_prompt=None, model="gpt-3.5-turbo"):
#     """
#     Send a query to ChatGPT 3.5 and get the response.
    
#     Parameters:
#     prompt (str): The query text to send to ChatGPT
#     api_key (str): Your OpenAI API key
#     system_prompt (str, optional): System instructions for ChatGPT
#     model (str): OpenAI model to use (default: gpt-3.5-turbo)
    
#     Returns:
#     str: The response from ChatGPT
#     """
#     headers = {
#         "Content-Type": "application/json",
#         "Authorization": f"Bearer {openai.api_key}"
#     }
    
#     # Start with an empty messages list
#     messages = []
    
#     # Add system prompt if provided
#     if system_prompt:
#         messages.append({"role": "system", "content": system_prompt})
    
#     # Add user prompt
#     messages.append({"role": "user", "content": prompt})
    
#     payload = {
#         "model": model,
#         "messages": messages,
#         "temperature": 0.2
#     }
    
#     try:
#         response = requests.post(
#             "https://api.openai.com/v1/chat/completions",
#             headers=headers,
#             data=json.dumps(payload)
#         )
        
#         # Check if the request was successful
#         response.raise_for_status()
        
#         # Parse the response
#         response_data = response.json()
        
#         # Extract the message content
#         return response_data["choices"][0]["message"]["content"]
    
#     except requests.exceptions.HTTPError as e:
#         return f"HTTP Error: {e}"
#     except requests.exceptions.ConnectionError:
#         return "Error: Connection failed"
#     except KeyError:
#         return f"Error: Unexpected response format: {response_data}"
#     except Exception as e:
#         return f"Error: {str(e)}"


def format_Prompt_to_LLM(doc_count, results, user_quary):
    prompt_to_llm = "\n".join([f"\n\n<CONTEXT_{i+1}>\n\n{doc['text']}\n\n</CONTEXT_{i+1}>" for i, doc in enumerate(results[:doc_count])])
    prompt_to_llm = prompt_to_llm + f"\n\nThis is user quary : {user_quary} "
    return prompt_to_llm


# def process_documents_and_save_embeddings(
#         file_root,
#         collection_name,
#         embedding_model,
#         chroma_client,
#         retrive_method: str = 'fileNames',
#         file_names: List[str] = None,
#         folder_name: str = None
#         ):
    
#     """
#     process documents and save embeddings.
    
#     Args:
#         file_root (str): root path in saved documents
#         collection_name (str): collection name 
#         embedding_model (str): embedding model name
#         chroma_client (ClientAPI): chroma DB client
#         retrive_method (str): file retrivel methode (fileNames/folderName)
#         file_names (List[str], optional): files name list, when use only retrive_methode="fileNames"
#         folder_name (str, optional): collection folder name, when use only retrive_methode="folderName"

#     Returns:
#     """
    
#     if retrive_method == "fileNames":
#         documents = get_documents_from_filenames(file_root, file_names, collection_name)
#     elif retrive_method == "folderName":
#         clear_chroma_database(folder_name)
#         documents = get_documents_from_folder(rf"{file_root}\{folder_name}")

#     # documents = get_documents_from_filenames("files", file_names, collection_name)
#     #documents = get_documents_from_folder("files\cars")
#     #documents = extract_documents_from_zip(zip_path)

#     #save_documents_as_text_files_in_folder(documents)

#     chunks = chunk_text(chunk_size=1000, chunk_overlap=100, documents=documents)
#     #save_chunks_as_text_files(chunks=chunks)
    
#     get_and_save_embedings(batch_size=10, chunks=chunks, embedding_model=embedding_model, chroma_client=chroma_client, collection_name=collection_name)



def remove_data_by_filenames(chroma_client, collection_name, file_name_list):
    """
    Remove documents from a Chroma collection based on a list of file names.
    
    Parameters:
    -----------
    chroma_client : ChromaClient
        The initialized Chroma client instance
    collection_name : str
        Name of the collection to remove data from
    file_name_list : list
        List of file names to remove from the collection
    
    Returns:
    --------
    int
        Number of documents removed
    """
    try:
        #file name re genarate avoiding extention (all are txt)
        file_name_list = [nn.split('.')[0]+'.txt' for nn in file_name_list]
        # Get the collection
        collection = chroma_client.get_collection(collection_name)
        print(f"Accessing collection: {collection_name}, {file_name_list} ,{type(file_name_list)}")
        
        # Get all documents with their metadata
        results = collection.get(include=["metadatas", "documents", "embeddings"])
        
        # Find IDs of documents that match the file names in file_name_list
        ids_to_remove = []
        for i, metadata in enumerate(results["metadatas"]):
            print(f"{i} : {metadata}")
            if metadata.get("source") in file_name_list:
                ids_to_remove.append(results["ids"][i])
        
        # Remove the documents if any were found
        if ids_to_remove:
            collection.delete(ids=ids_to_remove)
            print(f"Removed {len(ids_to_remove)} documents from {collection_name}")
            return len(ids_to_remove)
        else:
            print(f"No documents matching the provided file names were found")
            return 0
            
    except Exception as e:
        print(f"Error removing data from collection {collection_name}: {str(e)}")
        return 0